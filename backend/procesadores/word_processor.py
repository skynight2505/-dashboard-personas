from docx import Document
import logging

from procesadores.pdf_processor import extraer_personas_desde_texto

logger = logging.getLogger(__name__)


def procesar_word(ruta_archivo: str) -> list[dict]:
    try:
        doc = Document(ruta_archivo)
        texto_completo = "\n".join([p.text for p in doc.paragraphs])

        for tabla in doc.tables:
            for fila in tabla.rows:
                texto_fila = " | ".join([celda.text for celda in fila.cells])
                texto_completo += "\n" + texto_fila

        return extraer_personas_desde_texto(texto_completo)

    except Exception as e:
        logger.error(f"Error al procesar Word {ruta_archivo}: {e}")
        raise
